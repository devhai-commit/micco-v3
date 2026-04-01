from __future__ import annotations
from fastapi import APIRouter, Depends, HTTPException, Body, BackgroundTasks
from sqlalchemy import select, update, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.deps import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.document import Document, DocumentStatus
from app.models.knowledge_entry import KnowledgeEntry
from app.models.department import Department
from app.api_compat.utils import format_bytes_to_human, get_current_department_id
from app.api.documents import process_document_background, process_knowledge_background, UPLOAD_DIR
from docx import Document as DocxDocument
import aiofiles
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/approvals", tags=["Approvals"])


@router.get("/count")
async def pending_count(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        return {"count": 0}

    # Trưởng phòng chỉ thấy pending của phòng mình
    dept_filter_doc = True
    dept_filter_kn = True
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        dept_filter_doc = Document.department_id == current_user.department_id
        dept_filter_kn = KnowledgeEntry.department_id == current_user.department_id

    doc_stmt = select(func.count(Document.id)).where(
        Document.approval_status == "pending",
        dept_filter_doc,
    )
    doc_count = (await db.execute(doc_stmt)).scalar() or 0

    kn_stmt = select(func.count(KnowledgeEntry.id)).where(
        KnowledgeEntry.approval_status == "pending_approval",
        dept_filter_kn,
    )
    kn_count = (await db.execute(kn_stmt)).scalar() or 0

    total = doc_count + kn_count
    last_requester = None

    if total > 0:
        # Fetch the most recent pending item (doc or knowledge) scoped by department
        latest_doc_result = await db.execute(
            select(User.name, Document.created_at)
            .join(User, Document.uploader_id == User.id)
            .where(
                Document.approval_status == "pending",
                dept_filter_doc,
            )
            .order_by(Document.created_at.desc())
            .limit(1)
        )
        doc_row = latest_doc_result.first()

        latest_kn_result = await db.execute(
            select(User.name, KnowledgeEntry.created_at)
            .join(User, KnowledgeEntry.owner_id == User.id)
            .where(
                KnowledgeEntry.approval_status == "pending_approval",
                dept_filter_kn,
            )
            .order_by(KnowledgeEntry.created_at.desc())
            .limit(1)
        )
        kn_row = latest_kn_result.first()

        if doc_row and kn_row:
            last_requester = doc_row[0] if doc_row[1] > kn_row[1] else kn_row[0]
        elif doc_row:
            last_requester = doc_row[0]
        elif kn_row:
            last_requester = kn_row[0]

    return {"count": total, "last_requester": last_requester}



@router.get("/pending")
async def list_pending(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        return {"documents": [], "knowledge": []}

    # Trưởng phòng chỉ thấy pending của phòng mình
    dept_filter_doc = True
    dept_filter_kn = True
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        dept_filter_doc = Document.department_id == current_user.department_id
        dept_filter_kn = KnowledgeEntry.department_id == current_user.department_id

    # Join with User and Department to get names
    stmt = (
        select(Document, User.name, Department.name)
        .outerjoin(User, Document.uploader_id == User.id)
        .outerjoin(Department, Document.department_id == Department.id)
        .where(
            Document.approval_status == "pending",
            dept_filter_doc,
        )
        .order_by(Document.created_at.desc())
    )
    result = await db.execute(stmt)
    rows = result.all()

    docs = []
    for doc, uploader_name, dept_name in rows:
        docs.append({
            "id": doc.id,
            "name": doc.original_filename,
            "category": "Tài liệu",
            "owner": uploader_name or "Hệ thống",
            "department": dept_name or "Chung",
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "size": format_bytes_to_human(doc.file_size or 0),
            "visibility": doc.visibility,
            "file_type": doc.file_type.lower(),
        })

    # Fetch pending knowledge entries
    kn_stmt = (
        select(KnowledgeEntry, User.name, Department.name)
        .outerjoin(User, KnowledgeEntry.owner_id == User.id)
        .outerjoin(Department, KnowledgeEntry.department_id == Department.id)
        .where(
            KnowledgeEntry.approval_status == "pending_approval",
            dept_filter_kn,
        )
        .order_by(KnowledgeEntry.created_at.desc())
    )
    kn_result = await db.execute(kn_stmt)
    kn_rows = kn_result.all()

    knowledge_items = []
    for entry, owner_name, dept_name in kn_rows:
        knowledge_items.append({
            "id": entry.id,
            "title": entry.title,
            "content_text": entry.content_text,
            "category": entry.category,
            "owner": owner_name or "Hệ thống",
            "department": dept_name or "Chung",
            "created_at": entry.created_at.isoformat() if entry.created_at else None,
            "visibility": entry.visibility,
            "tags": entry.tags
        })

    return {"documents": docs, "knowledge": knowledge_items}


@router.post("/documents/{doc_id}/approve")
async def approve_document(
    doc_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Trưởng phòng chỉ được phê duyệt tài liệu của phòng mình
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        if doc.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể phê duyệt tài liệu của phòng ban khác"
            )

    doc.approval_status = "approved"
    doc.status = DocumentStatus.PROCESSING  # Set to processing as we start indexing
    await db.commit()

    # Trigger background parsing & indexing (NexusRAG style)
    from app.core.config import settings
    file_path = str(settings.BASE_DIR / "uploads" / doc.filename)
    background_tasks.add_task(process_document_background, doc.id, file_path, doc.workspace_id)

    return {"message": "Đã phê duyệt và đang bắt đầu xử lý", "id": doc_id}


@router.post("/documents/{doc_id}/reject")
async def reject_document(
    doc_id: int,
    note: str = Body(None, embed=True),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Trưởng phòng chỉ được từ chối tài liệu của phòng mình
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        if doc.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể từ chối tài liệu của phòng ban khác"
            )

    doc.approval_status = "rejected"
    doc.approval_note = note
    await db.commit()
@router.get("/documents/{doc_id}/status")
async def get_document_status(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Return current processing status of an approved document."""
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return {
        "id": doc.id,
        "status": doc.status.value if hasattr(doc.status, "value") else doc.status,
        "approval_status": doc.approval_status,
        "chunk_count": doc.chunk_count,
        "error_message": doc.error_message,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
    }


@router.get("/documents/{doc_id}/preview")
async def preview_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    doc = (await db.execute(select(Document).where(Document.id == doc_id))).scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = UPLOAD_DIR / doc.filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = doc.file_type.lower()
    text_content = ""

    try:
        if ext == "docx":
            # Quick sync read for docx (safe for small snippets)
            d = DocxDocument(str(file_path))
            # Get first 30 paragraphs
            text_content = "\n".join([p.text for p in d.paragraphs[:30]])
            if len(d.paragraphs) > 30:
                text_content += "\n\n...(Còn tiếp)..."
        elif ext in ("txt", "md"):
            async with aiofiles.open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                text_content = await f.read(5000) # first 5k chars
                if len(text_content) == 5000:
                    text_content += "\n\n...(Còn tiếp)..."
        else:
            return {"supported": False, "message": "Preview not supported for this type"}
            
        return {"supported": True, "content": text_content, "file_type": ext}
    except Exception as e:
        logger.error(f"Error previewing file {file_path}: {str(e)}")
        return {"supported": False, "message": f"Error loading preview: {str(e)}"}


@router.post("/knowledge/{entry_id}/approve")
async def approve_knowledge(
    entry_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    entry = (await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id)
    )).scalar_one_or_none()

    if not entry:
        raise HTTPException(status_code=404, detail="Knowledge entry not found")

    # Trưởng phòng chỉ được phê duyệt tri thức của phòng mình
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        if entry.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể phê duyệt tri thức của phòng ban khác"
            )

    entry.approval_status = "approved"
    entry.status = "Active"
    entry.ingest_status = "processing"
    await db.commit()

    # Trigger background indexing
    background_tasks.add_task(process_knowledge_background, entry_id, 1)  # Default workspace 1

    return {"message": "Đã phê duyệt tri thức", "id": entry_id}

@router.post("/knowledge/{entry_id}/reject")
async def reject_knowledge(
    entry_id: int,
    note: str = Body(None, embed=True),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if current_user.role not in ("Admin", "Trưởng phòng"):
        raise HTTPException(status_code=403, detail="Permission denied")

    entry = (await db.execute(
        select(KnowledgeEntry).where(KnowledgeEntry.id == entry_id)
    )).scalar_one_or_none()

    if not entry:
        raise HTTPException(status_code=404, detail="Knowledge entry not found")

    # Trưởng phòng chỉ được từ chối tri thức của phòng mình
    if current_user.role == "Trưởng phòng" and current_user.department_id is not None:
        if entry.department_id != current_user.department_id:
            raise HTTPException(
                status_code=403,
                detail="Không thể từ chối tri thức của phòng ban khác"
            )

    entry.approval_status = "rejected"
    entry.approval_note = note
    entry.status = "Draft"
    await db.commit()
    return {"message": "Đã từ chối tri thức", "id": entry_id}
